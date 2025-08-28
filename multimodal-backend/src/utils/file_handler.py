"""
File Upload and Processing Utilities
Handles multimodal file uploads, processing, and storage
"""

import os
import uuid
import magic
import shutil
from pathlib import Path
from typing import List, Dict, Any, Optional
from PIL import Image
import PyPDF2
from docx import Document
import pandas as pd
import cv2
import pytesseract
import base64
from fastapi import UploadFile, HTTPException


class FileHandler:
    """Handles file uploads, validation, and processing"""
    
    ALLOWED_EXTENSIONS = {
        'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'],
        'document': ['.pdf', '.docx', '.doc', '.txt'],
        'spreadsheet': ['.xlsx', '.xls', '.csv']
    }
    
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        
        # Create subdirectories
        (self.upload_dir / "images").mkdir(exist_ok=True)
        (self.upload_dir / "documents").mkdir(exist_ok=True)
        (self.upload_dir / "processed").mkdir(exist_ok=True)
    
    def validate_file(self, file: UploadFile) -> Dict[str, Any]:
        """Validate uploaded file"""
        if file.size > self.MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail="File too large")
        
        # Get file extension
        file_ext = Path(file.filename).suffix.lower()
        
        # Determine file type
        file_type = None
        for category, extensions in self.ALLOWED_EXTENSIONS.items():
            if file_ext in extensions:
                file_type = category
                break
        
        if not file_type:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type: {file_ext}"
            )
        
        return {
            "type": file_type,
            "extension": file_ext,
            "filename": file.filename,
            "size": file.size
        }
    
    async def save_file(self, file: UploadFile) -> Dict[str, Any]:
        """Save uploaded file and return file info"""
        file_info = self.validate_file(file)
        
        # Generate unique filename
        unique_id = str(uuid.uuid4())
        safe_filename = f"{unique_id}_{file.filename}"
        
        # Determine save path based on file type
        if file_info["type"] == "image":
            file_path = self.upload_dir / "images" / safe_filename
        else:
            file_path = self.upload_dir / "documents" / safe_filename
        
        # Save file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Add path info to file_info
        file_info.update({
            "id": unique_id,
            "saved_filename": safe_filename,
            "file_path": str(file_path),
            "relative_path": str(file_path.relative_to(self.upload_dir))
        })
        
        return file_info
    
    def process_image(self, file_path: str) -> Dict[str, Any]:
        """Process image file - extract metadata, perform OCR"""
        try:
            # Open image
            image = Image.open(file_path)
            
            # Basic metadata
            metadata = {
                "size": image.size,
                "mode": image.mode,
                "format": image.format
            }
            
            # Perform OCR
            ocr_text = ""
            try:
                # Convert PIL image to OpenCV format for OCR
                cv_image = cv2.imread(file_path)
                ocr_text = pytesseract.image_to_string(cv_image)
            except Exception as e:
                print(f"OCR failed: {e}")
            
            # Convert to base64 for API responses
            with open(file_path, "rb") as img_file:
                img_base64 = base64.b64encode(img_file.read()).decode()
            
            return {
                "metadata": metadata,
                "ocr_text": ocr_text.strip(),
                "base64": img_base64,
                "processed": True
            }
            
        except Exception as e:
            return {
                "error": str(e),
                "processed": False
            }
    
    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF file - extract text and metadata"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                metadata = {
                    "pages": len(reader.pages),
                    "title": reader.metadata.get('/Title', '') if reader.metadata else '',
                    "author": reader.metadata.get('/Author', '') if reader.metadata else ''
                }
                
                # Extract text
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    "metadata": metadata,
                    "text": text.strip(),
                    "processed": True
                }
                
        except Exception as e:
            return {
                "error": str(e),
                "processed": False
            }
    
    def process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process Word document"""
        try:
            doc = Document(file_path)
            
            # Extract text
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Basic metadata
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "has_tables": len(doc.tables) > 0,
                "tables_count": len(doc.tables)
            }
            
            return {
                "metadata": metadata,
                "text": text.strip(),
                "processed": True
            }
            
        except Exception as e:
            return {
                "error": str(e),
                "processed": False
            }
    
    def process_spreadsheet(self, file_path: str) -> Dict[str, Any]:
        """Process Excel/CSV file"""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Basic analysis
            metadata = {
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist(),
                "dtypes": df.dtypes.to_dict()
            }
            
            # Convert to text representation
            text = df.to_string()
            
            return {
                "metadata": metadata,
                "text": text,
                "dataframe_info": df.describe().to_dict(),
                "processed": True
            }
            
        except Exception as e:
            return {
                "error": str(e),
                "processed": False
            }
    
    def process_file(self, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Process file based on type"""
        file_path = file_info["file_path"]
        file_type = file_info["type"]
        
        if file_type == "image":
            processing_result = self.process_image(file_path)
        elif file_info["extension"] == ".pdf":
            processing_result = self.process_pdf(file_path)
        elif file_info["extension"] in [".docx", ".doc"]:
            processing_result = self.process_docx(file_path)
        elif file_info["extension"] in [".xlsx", ".xls", ".csv"]:
            processing_result = self.process_spreadsheet(file_path)
        elif file_info["extension"] == ".txt":
            # Simple text file
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            processing_result = {
                "text": text,
                "metadata": {"size": len(text)},
                "processed": True
            }
        else:
            processing_result = {
                "error": "Unsupported file type for processing",
                "processed": False
            }
        
        # Combine file info with processing result
        result = {**file_info, **processing_result}
        return result
    
    def cleanup_file(self, file_path: str) -> bool:
        """Remove uploaded file"""
        try:
            os.remove(file_path)
            return True
        except Exception as e:
            print(f"Failed to cleanup file {file_path}: {e}")
            return False